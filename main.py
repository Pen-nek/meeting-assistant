"""
meeting-assistant · main.py
실행: uvicorn main:app --reload
"""
import asyncio, io, json, os, re, uuid
from contextlib import asynccontextmanager
from datetime import datetime
from typing import List, Optional

import httpx
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel

load_dotenv()
GLADIA_KEY = os.getenv("GLADIA_API_KEY", "")
GROQ_KEY   = os.getenv("GROQ_API_KEY",   "")
GROQ_MODEL      = "llama-3.1-8b-instant"
GROQ_FALLBACK   = "llama-3.3-70b-versatile"
MAX_CHARS       = 8000

http: httpx.AsyncClient | None = None
_jobs: dict = {}

@asynccontextmanager
async def lifespan(app: FastAPI):
    global http
    http = httpx.AsyncClient(timeout=120)
    yield
    await http.aclose()

app = FastAPI(lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


# ── 모델 ─────────────────────────────────────────
class Todo(BaseModel):
    task:     str
    owner:    Optional[str] = None
    speaker:  Optional[int] = None
    priority: Optional[str] = "보통"


# ── 프롬프트 ──────────────────────────────────────
_SYSTEM_KO = (
    "You are a Korean meeting assistant. "
    "Write ALL text values in Korean (한글) only. "
    "Never use Chinese characters, Japanese, or English in text values. "
    "JSON keys stay in English; all values must be Korean."
)

_MINUTES = """아래 규칙을 반드시 따르세요:
1. 반드시 한국어로만 작성하세요. 외국 문자 사용 금지.
2. 불분명한 내용은 추측하지 말고 생략하세요.
3. 결정된 사항은 대화에서 명확히 결정된 것만 적으세요. 없으면 "없음".
4. 문체는 사무적 문어체로 통일하세요. (~됨, ~결정됨, ~논의됨, ~예정임 등 명사형 또는 완료형 종결어미 사용. 존댓말·구어체 사용 금지)
5. 다음 형식을 그대로 사용하세요:

# 회의 요약
(2~3문장)

## 주요 논의사항
(간결하게)

## 결정된 사항
(구체적으로)

## 특이사항
(없으면 생략)"""

_TODO = """대화를 정확히 분석하여 확정된 액션 아이템만 추출하세요.

추출 기준 (모두 충족해야 포함):
1. 특정 인물이 "~하겠다", "~할게요", "~하기로 했다" 등 명시적으로 약속하거나 확정한 것
2. 단순 제안·논의·가능성 언급은 제외 (예: "~하면 좋겠다", "~어떨까요" 등)
3. task는 주어 없이 동사로 시작하는 개조식 한국어 (예: "~하기", "~공유" 등)
4. 불명확한 내용은 추측하지 말고 생략

필드 규칙:
- owner: 대화에서 직접 언급된 담당자 이름. 불명확하면 null
- speaker: 담당자의 발화자 번호(정수). 불명확하면 null
- priority: 마감 언급·긴급도 기준. 기본값 "보통"
- 반드시 한국어로만 작성. 외국 문자 절대 금지
- task 내용은 반드시 실제 대화에서 언급된 내용만 사용할 것. 아래 JSON 형식의 key 이름이나 예시 문자열을 task 값으로 절대 사용하지 말 것

JSON 배열만 응답 (설명 없이):
[{"task":"<실제 대화 내용 기반>","owner":"<담당자 또는 null>","speaker":0,"priority":"높음|보통|낮음"}]"""


# ── 유틸 ─────────────────────────────────────────
def _truncate(text: str) -> str:
    if len(text) <= MAX_CHARS:
        return text
    head = int(MAX_CHARS * 0.7)
    return text[:head] + "\n\n...(중략)...\n\n" + text[-(MAX_CHARS - head):]

def _strip_cjk(text: str) -> str:
    return re.sub(r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]+', '', text)

def _parse_json(raw: str, fallback):
    try:
        return json.loads(raw.replace("```json", "").replace("```", "").strip())
    except Exception:
        return fallback

def _parse_todos(raw: str) -> List[Todo]:
    try:
        return [Todo(**t) for t in _parse_json(raw, [])]
    except Exception:
        return []

def _check_keys():
    if not GLADIA_KEY or not GROQ_KEY:
        raise HTTPException(500, ".env 파일에 API 키를 설정해주세요.")


# ── Gladia: 음성 인식 + 화자 분리 ───────────────
async def transcribe(audio: bytes, n_speakers: int, lang: str) -> list:
    headers   = {"x-gladia-key": GLADIA_KEY}
    lang_code = {"ko": "ko", "ja": "ja", "zh": "zh"}.get(lang, "en")

    up = await http.post(
        "https://api.gladia.io/v2/upload",
        headers=headers,
        files={"audio": ("audio", audio, "audio/mpeg")},
    )
    up.raise_for_status()

    tr = await http.post(
        "https://api.gladia.io/v2/transcription",
        headers={**headers, "Content-Type": "application/json"},
        json={
            "audio_url": up.json()["audio_url"],
            "language": lang_code,
            "diarization": True,
            "diarization_config": {"number_of_speakers": n_speakers, "max_speakers": n_speakers},
        },
    )
    tr.raise_for_status()
    result_url = tr.json()["result_url"]

    for i in range(120):
        await asyncio.sleep(3)
        data = (await http.get(result_url, headers=headers)).json()
        print(f"[Gladia poll #{i}] status={data.get('status')}")
        if data["status"] == "done":
            utterances = data["result"]["transcription"]["utterances"]
            for u in utterances:
                u.setdefault("speaker", 0)
            return utterances
        if data["status"] in ("error", "failed"):
            raise HTTPException(500, f"음성 인식 실패: {data.get('error_code')}")
    raise HTTPException(500, "음성 인식 타임아웃")


# ── Groq: LLM 호출 ───────────────────────────────
async def groq(prompt: str, max_tokens: int = 2000, retries: int = 4) -> str:
    delay = 5.0
    for attempt in range(retries):
        model = GROQ_MODEL if attempt < 2 else GROQ_FALLBACK
        res = await http.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_KEY}", "Content-Type": "application/json"},
            json={
                "model": model,
                "max_tokens": max_tokens,
                "messages": [
                    {"role": "system", "content": _SYSTEM_KO},
                    {"role": "user",   "content": prompt},
                ],
            },
        )
        if res.status_code == 429:
            wait = float(res.headers.get("retry-after", delay))
            print(f"[Groq 429] model={model} retry_after={wait}s")
            if wait > 60:
                raise HTTPException(429, f"Groq API 사용량 초과 ({int(wait)}초 대기 필요). 잠시 후 다시 시도해주세요.")
            await asyncio.sleep(wait)
            delay = min(delay * 2, 60)
            continue
        res.raise_for_status()
        return _strip_cjk(res.json()["choices"][0]["message"]["content"])
    raise HTTPException(429, "Groq API 요청 한도 초과. 잠시 후 다시 시도해주세요.")


# ── Word 생성 ─────────────────────────────────────
def build_docx(minutes: str, todos: List[Todo], transcript: str = "") -> bytes:
    doc = Document()
    title = doc.add_heading("회의록", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date = doc.add_paragraph(datetime.now().strftime("%Y년 %m월 %d일")); date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.size = Pt(10); date.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x88)
    doc.add_paragraph()

    if minutes:
        doc.add_heading("📋 회의록", 1)
        for line in minutes.split("\n"):
            if line.startswith(("## ", "# ")): doc.add_heading(line.lstrip("# "), 2)
            elif line.strip(): doc.add_paragraph(line)

    if todos:
        doc.add_page_break()
        doc.add_heading("✅ TO DO", 1)
        tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
        for i, h in enumerate(["할 일", "담당자", "우선순위", "완료"]):
            c = tbl.rows[0].cells[i]; c.text = h; c.paragraphs[0].runs[0].font.bold = True
        for t in todos:
            r = tbl.add_row().cells
            r[0].text = t.task; r[1].text = t.owner or "-"; r[2].text = t.priority or "보통"; r[3].text = "☐"

    if transcript:
        doc.add_page_break()
        doc.add_heading("📝 전체 대화", 1)
        for line in transcript.split("\n"):
            if line.strip(): doc.add_paragraph(line)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ── 엔드포인트 ────────────────────────────────────
@app.post("/analyze")
async def analyze_endpoint(
        file: UploadFile = File(...),
        speaker_count: int = Form(4),
        language: str = Form("ko"),
):
    _check_keys()
    audio  = await file.read()
    job_id = str(uuid.uuid4())

    async def stream():
        def sse(event: str, data) -> str:
            return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"
        try:
            yield sse("progress", {"step": 2})
            utterances   = await transcribe(audio, speaker_count, language)
            full_text    = "\n".join(f"[발화자 {u['speaker']}] {u['text']}" for u in utterances)
            speaker_list = ", ".join(str(s) for s in sorted({u["speaker"] for u in utterances}))

            yield sse("progress", {"step": 3})
            minutes_raw = await groq(f"다음은 회의 대화 내용입니다.\n\n{_truncate(full_text)}\n\n{_MINUTES}")

            yield sse("progress", {"step": 4})
            await asyncio.sleep(1)
            todo_raw = await groq(
                f"다음 회의 대화에서 액션 아이템을 추출하세요.\n\n{_truncate(full_text)}\n\n발화자 번호: {speaker_list}\n\n{_TODO}",
                max_tokens=1000,
            )
            todos = _parse_todos(todo_raw)
            _jobs[job_id] = {
                "utterances": utterances,
                "full_text":  full_text,
                "minutes":    minutes_raw,
                "todos":      [t.dict() for t in todos],
            }
            yield sse("done", {"job_id": job_id})
        except HTTPException as e:
            yield sse("error", {"detail": e.detail})
        except Exception as e:
            yield sse("error", {"detail": str(e)})

    return StreamingResponse(stream(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.get("/result/{job_id}")
async def get_result(job_id: str):
    result = _jobs.pop(job_id, None)
    if result is None:
        raise HTTPException(404, "결과를 찾을 수 없습니다.")
    return result


@app.post("/regenerate")
async def regenerate_endpoint(full_text: str = Form(...), speaker_count: int = Form(4)):
    _check_keys()
    lines = [l for l in full_text.split("\n") if l.strip()]
    speaker_list = ", ".join(str(s) for s in sorted({
        int(l.split("]")[0].replace("[발화자 ", ""))
        for l in lines if l.startswith("[발화자 ")
    }))
    t = _truncate(full_text)
    minutes_raw = await groq(f"다음은 회의 대화 내용입니다.\n\n{t}\n\n{_MINUTES}")
    await asyncio.sleep(1)
    todo_raw    = await groq(
        f"다음 회의 대화에서 액션 아이템을 추출하세요.\n\n{t}\n\n발화자 번호: {speaker_list}\n\n{_TODO}",
        max_tokens=1000,
    )
    todos = _parse_todos(todo_raw)
    return {"minutes": minutes_raw, "todos": [t.dict() for t in todos]}


@app.post("/generate-docx")
async def docx_endpoint(
        minutes:    str = Form(""),
        todos_json: str = Form("[]"),
        transcript: str = Form(""),
):
    todos = _parse_todos(todos_json)
    return StreamingResponse(
        io.BytesIO(build_docx(minutes, todos, transcript)),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=meeting_minutes.docx"},
    )


@app.get("/")
def root():
    return {"status": "meeting-assistant running"}